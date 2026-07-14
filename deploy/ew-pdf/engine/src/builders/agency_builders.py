"""DOCX builders for agency-specific bulletins (MoW, GST, MoH, MoA, NEMC)."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .base_builder import BaseBulletinBuilder
from .styles import (
    FONT_PRIMARY, COLOR_BLACK, COLOR_WHITE,
    SIZE_TITLE, SIZE_SUBTITLE, SIZE_BODY, SIZE_BODY_SMALL, SIZE_SMALL,
    SIZE_LABEL, SIZE_FOOTER,
    HEX_YELLOW_ADVISORY, HEX_ORANGE_WARNING, HEX_RED_MAJOR, HEX_WHITE,
    MAP_DAY1_WIDTH, MAP_DAY1_HEIGHT,
)
from .table_helpers import (
    set_cell_shading, set_cell_borders, remove_all_borders,
    set_table_fixed_layout, set_table_col_widths, set_table_width,
    set_cell_vertical_alignment, set_cell_margins,
    set_paragraph_shading,
)
from .header_helpers import _add_run, _try_add_image, ASSETS_DIR

from ..models.agency import (
    MoWBulletin, GSTBulletin, MoHBulletin, MoABulletin, NEMCBulletin,
)
from ..models.common import AlertLevel


_ALERT_HEX = {
    AlertLevel.ADVISORY: HEX_YELLOW_ADVISORY,
    AlertLevel.WARNING: HEX_ORANGE_WARNING,
    AlertLevel.MAJOR_WARNING: HEX_RED_MAJOR,
    AlertLevel.NO_WARNING: HEX_WHITE,
}

_ALERT_LABEL = {
    AlertLevel.ADVISORY: "ADVISORY",
    AlertLevel.WARNING: "WARNING",
    AlertLevel.MAJOR_WARNING: "MAJOR WARNING",
    AlertLevel.NO_WARNING: "NO WARNING",
}

_ALERT_TEXT_COLOR = {
    AlertLevel.ADVISORY: COLOR_BLACK,
    AlertLevel.WARNING: COLOR_BLACK,
    AlertLevel.MAJOR_WARNING: COLOR_BLACK,
    AlertLevel.NO_WARNING: COLOR_BLACK,
}


def _build_agency_header(doc, title: str, subtitle: str, date_str: str, time_str: str):
    """Build a standard agency header: [Coat of Arms | Title + subtitle + date]."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    coat_w, title_w = 1418, 9072  # ~2.5cm, ~16cm
    set_table_fixed_layout(table)
    set_table_col_widths(table, [coat_w, title_w])
    set_table_width(table, coat_w + title_w)

    for cell in table.rows[0].cells:
        remove_all_borders(cell)
        set_cell_vertical_alignment(cell, "center")

    # Coat of Arms
    coat_path = ASSETS_DIR / "logos" / "coat_of_arms.jpeg"
    _try_add_image(table.cell(0, 0), coat_path, Cm(1.8), Cm(1.8))

    # Title block
    cell = table.cell(0, 1)
    cell.paragraphs[0].clear()

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, "UNITED REPUBLIC OF TANZANIA", size=SIZE_SUBTITLE, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, title, size=SIZE_TITLE, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, subtitle, size=SIZE_BODY, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, f"Issued: {date_str}  {time_str}", size=SIZE_BODY_SMALL)

    return table


def _add_alert_badge(doc, level: AlertLevel):
    """Add a colored alert level badge as a paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    hex_color = _ALERT_HEX.get(level, HEX_YELLOW_ADVISORY)
    set_paragraph_shading(p, hex_color)
    text_color = _ALERT_TEXT_COLOR.get(level, COLOR_BLACK)
    _add_run(p, f"  {_ALERT_LABEL.get(level, 'ADVISORY')}  ",
             size=SIZE_BODY, bold=True, color=text_color)
    return p


def _add_field_row(doc, label: str, value: str):
    """Add a label: value row."""
    if not value:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    _add_run(p, f"{label}: ", size=SIZE_BODY_SMALL, bold=True)
    _add_run(p, value, size=SIZE_BODY_SMALL)


def _add_section_heading(doc, text: str):
    """Add a section heading with blue background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_shading(p, "D6E8F0")
    _add_run(p, f"  {text}", size=SIZE_BODY, bold=True)


def _add_footer(doc, agency: str, date_str: str):
    """Add a simple footer."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"{agency} — Issued {date_str}", size=SIZE_FOOTER)
    _add_run(p, "\nUnited Republic of Tanzania", size=SIZE_FOOTER)


# ===================================================================
# MoW Builder
# ===================================================================

class MoWBuilder(BaseBulletinBuilder):
    """Builds MoW Flood Risk Assessment bulletin."""

    def __init__(self, data: MoWBulletin, **kwargs):
        super().__init__(data, **kwargs)
        self._generated_maps = {}

    def build(self) -> Document:
        bulletin = self.data
        date_str = bulletin.issue_date.strftime("%d %B %Y")
        time_str = bulletin.issue_time.strftime("%H:%M")

        _build_agency_header(
            self.doc,
            "MINISTRY OF WATER",
            "Flood Risk Assessment — Three Days Forecast",
            date_str, time_str,
        )

        for i, day in enumerate(bulletin.days):
            _add_section_heading(
                self.doc,
                f"Day {i+1} — {day.forecast_date.strftime('%A %d %B %Y')}",
            )

            if not day.assessments:
                p = self.doc.add_paragraph()
                _add_run(p, "No flood risk assessment for this day.",
                         size=SIZE_BODY_SMALL)
                continue

            # Map
            map_path = self._generated_maps.get(i) or \
                       (day.map_image.file_path if day.map_image else None)
            if map_path:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_or_placeholder(p, map_path, Cm(8), Cm(6.5))

            for j, assessment in enumerate(day.assessments):
                _add_alert_badge(self.doc, assessment.alert_level)
                basins = ", ".join(assessment.catchment_basins) if assessment.catchment_basins else ""
                _add_field_row(self.doc, "Catchment Basins", basins)
                regions = ", ".join(assessment.regions) if assessment.regions else ""
                _add_field_row(self.doc, "Affected Regions", regions)
                districts = ", ".join(assessment.districts) if assessment.districts else ""
                _add_field_row(self.doc, "Affected Districts", districts)
                _add_field_row(self.doc, "Description", assessment.description)
                _add_field_row(self.doc, "Likelihood", assessment.likelihood)
                _add_field_row(self.doc, "Impact", assessment.impact)
                if assessment.impacts_expected:
                    _add_field_row(self.doc, "Impacts Expected", assessment.impacts_expected)

            if i < len(bulletin.days) - 1:
                self.doc.add_paragraph()  # spacer

        _add_footer(self.doc, "Ministry of Water", date_str)
        return self.doc


# ===================================================================
# GST Builder
# ===================================================================

class GSTBuilder(BaseBulletinBuilder):
    """Builds GST Geological Hazard bulletin."""

    def __init__(self, data: GSTBulletin, **kwargs):
        super().__init__(data, **kwargs)
        self._generated_maps = {}

    def build(self) -> Document:
        bulletin = self.data
        date_str = bulletin.issue_date.strftime("%d %B %Y")
        time_str = bulletin.issue_time.strftime("%H:%M")

        _build_agency_header(
            self.doc,
            "GEOLOGICAL SURVEY OF TANZANIA",
            "Geological Hazard Assessment Bulletin",
            date_str, time_str,
        )

        for i, event in enumerate(bulletin.events):
            etype = event.event_type.replace("_", " ").title()
            _add_section_heading(self.doc, f"Event {i+1}: {etype}")
            _add_alert_badge(self.doc, event.alert_level)

            # Type-specific fields
            if event.event_type == "EARTHQUAKE":
                if event.magnitude is not None:
                    _add_field_row(self.doc, "Magnitude", str(event.magnitude))
                if event.depth_km is not None:
                    _add_field_row(self.doc, "Depth (km)", str(event.depth_km))
                if event.severity:
                    _add_field_row(self.doc, "Severity", event.severity)
            elif event.event_type == "VOLCANO":
                if event.volcanic_hazard_index:
                    _add_field_row(self.doc, "Volcanic Hazard Index",
                                   event.volcanic_hazard_index)
                if event.activity_type:
                    _add_field_row(self.doc, "Activity Type", event.activity_type)

            regions = ", ".join(event.regions) if event.regions else ""
            _add_field_row(self.doc, "Affected Regions", regions)
            districts = ", ".join(event.districts) if event.districts else ""
            _add_field_row(self.doc, "Affected Districts", districts)
            _add_field_row(self.doc, "Description", event.description)
            _add_field_row(self.doc, "Likelihood", event.likelihood)
            _add_field_row(self.doc, "Impact", event.impact)
            if event.impacts_expected:
                _add_field_row(self.doc, "Impacts Expected", event.impacts_expected)

            # Map
            map_path = self._generated_maps.get(i)
            if map_path:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_or_placeholder(p, map_path, Cm(8), Cm(6.5))

        _add_footer(self.doc, "Geological Survey of Tanzania", date_str)
        return self.doc


# ===================================================================
# MoH Builder
# ===================================================================

class MoHBuilder(BaseBulletinBuilder):
    """Builds MoH Disease Outbreak bulletin."""

    def __init__(self, data: MoHBulletin, **kwargs):
        super().__init__(data, **kwargs)
        self._generated_maps = {}

    def build(self) -> Document:
        bulletin = self.data
        date_str = bulletin.issue_date.strftime("%d %B %Y")
        time_str = bulletin.issue_time.strftime("%H:%M")

        _build_agency_header(
            self.doc,
            "MINISTRY OF HEALTH",
            "Disease Outbreak Monitoring Bulletin",
            date_str, time_str,
        )

        for i, outbreak in enumerate(bulletin.outbreaks):
            _add_section_heading(
                self.doc,
                f"Outbreak {i+1}: {outbreak.disease_type}",
            )
            _add_alert_badge(self.doc, outbreak.alert_level)

            # Epi data table
            table = self.doc.add_table(rows=1, cols=3)
            table.autofit = True
            for cell in table.rows[0].cells:
                remove_all_borders(cell)
            cells = table.rows[0].cells
            for idx, (lbl, val) in enumerate([
                ("Confirmed Cases", str(outbreak.confirmed_cases)),
                ("Deaths", str(outbreak.deaths)),
                ("Trend", outbreak.trend),
            ]):
                p = cells[idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, lbl + "\n", size=SIZE_SMALL, bold=True)
                _add_run(p, val, size=SIZE_TITLE, bold=True)

            regions = ", ".join(outbreak.regions) if outbreak.regions else ""
            _add_field_row(self.doc, "Affected Regions", regions)
            districts = ", ".join(outbreak.districts) if outbreak.districts else ""
            _add_field_row(self.doc, "Affected Districts", districts)
            _add_field_row(self.doc, "Situation Summary", outbreak.situation_summary)
            _add_field_row(self.doc, "Response Actions", outbreak.response_actions)
            _add_field_row(self.doc, "Spread Likelihood", outbreak.likelihood)
            _add_field_row(self.doc, "Health Impact", outbreak.impact)

            # Map
            map_path = self._generated_maps.get(i)
            if map_path:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_or_placeholder(p, map_path, Cm(8), Cm(6.5))

        _add_footer(self.doc, "Ministry of Health", date_str)
        return self.doc


# ===================================================================
# MoA Builder
# ===================================================================

class MoABuilder(BaseBulletinBuilder):
    """Builds MoA Agricultural Hazard bulletin."""

    def __init__(self, data: MoABulletin, **kwargs):
        super().__init__(data, **kwargs)
        self._generated_maps = {}

    def build(self) -> Document:
        bulletin = self.data
        date_str = bulletin.issue_date.strftime("%d %B %Y")
        time_str = bulletin.issue_time.strftime("%H:%M")

        _build_agency_header(
            self.doc,
            "MINISTRY OF AGRICULTURE",
            f"Agricultural Hazard Assessment — {bulletin.report_period} Report",
            date_str, time_str,
        )

        for i, assessment in enumerate(bulletin.assessments):
            htype = assessment.hazard_type.replace("_", " ").title()
            _add_section_heading(self.doc, f"Assessment {i+1}: {htype}")
            _add_alert_badge(self.doc, assessment.alert_level)

            if assessment.drought_severity:
                _add_field_row(self.doc, "Drought Severity", assessment.drought_severity)
            if assessment.rainfall_pct is not None:
                _add_field_row(self.doc, "Rainfall (% of normal)",
                               f"{assessment.rainfall_pct:.0f}%")
            if assessment.ndvi is not None:
                # ndvi is numeric (e.g. 0.31); _add_field_row writes a string run, so format it.
                _add_field_row(self.doc, "NDVI", f"{assessment.ndvi:.2f}")
            if assessment.affected_sectors:
                _add_field_row(self.doc, "Affected Sectors",
                               ", ".join(assessment.affected_sectors))

            regions = ", ".join(assessment.regions) if assessment.regions else ""
            _add_field_row(self.doc, "Affected Regions", regions)
            districts = ", ".join(assessment.districts) if assessment.districts else ""
            _add_field_row(self.doc, "Affected Districts", districts)
            _add_field_row(self.doc, "Situation Summary", assessment.situation_summary)
            _add_field_row(self.doc, "Recommended Actions", assessment.recommended_actions)
            _add_field_row(self.doc, "Worsening Likelihood", assessment.likelihood)
            _add_field_row(self.doc, "Food Security Impact", assessment.impact)

            # Map
            map_path = self._generated_maps.get(i)
            if map_path:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_or_placeholder(p, map_path, Cm(8), Cm(6.5))

        _add_footer(self.doc, "Ministry of Agriculture", date_str)
        return self.doc


# ===================================================================
# NEMC Builder
# ===================================================================

class NEMCBuilder(BaseBulletinBuilder):
    """Builds NEMC Environmental Hazard bulletin."""

    def __init__(self, data: NEMCBulletin, **kwargs):
        super().__init__(data, **kwargs)
        self._generated_maps = {}

    def build(self) -> Document:
        bulletin = self.data
        date_str = bulletin.issue_date.strftime("%d %B %Y")
        time_str = bulletin.issue_time.strftime("%H:%M")

        _build_agency_header(
            self.doc,
            "NATIONAL ENVIRONMENT MANAGEMENT COUNCIL",
            "Environmental Hazard Monitoring Bulletin",
            date_str, time_str,
        )

        for i, event in enumerate(bulletin.events):
            htype = event.hazard_type.replace("_", " ").title()
            _add_section_heading(self.doc, f"Event {i+1}: {htype}")
            _add_alert_badge(self.doc, event.alert_level)

            if event.pollution_source:
                _add_field_row(self.doc, "Pollution Source", event.pollution_source)
            if event.aqi_level:
                _add_field_row(self.doc, "AQI Level", event.aqi_level)
            if event.aqi_value is not None:
                _add_field_row(self.doc, "AQI Value", str(event.aqi_value))
            if event.key_pollutants:
                _add_field_row(self.doc, "Key Pollutants",
                               ", ".join(event.key_pollutants))

            regions = ", ".join(event.regions) if event.regions else ""
            _add_field_row(self.doc, "Affected Regions", regions)
            districts = ", ".join(event.districts) if event.districts else ""
            _add_field_row(self.doc, "Affected Districts", districts)
            _add_field_row(self.doc, "Situation Summary", event.situation_summary)
            _add_field_row(self.doc, "Health Advisory", event.health_advisory)
            _add_field_row(self.doc, "Persistence Likelihood", event.likelihood)
            _add_field_row(self.doc, "Health Impact", event.impact)

            # Map
            map_path = self._generated_maps.get(i)
            if map_path:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_or_placeholder(p, map_path, Cm(8), Cm(6.5))

        _add_footer(self.doc, "National Environment Management Council", date_str)
        return self.doc
