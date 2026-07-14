"""Builder for 722E_4 Five Days Severe Weather Impact-Based Forecast documents."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base_builder import BaseBulletinBuilder
from .styles import (
    COLOR_BLUE_TMA, COLOR_BLACK, COLOR_WHITE,
    HEX_YELLOW_ADVISORY, HEX_ORANGE_WARNING, HEX_RED_MAJOR,
    HEX_HIGHLIGHT_ADVISORY, HEX_HIGHLIGHT_WARNING, HEX_HIGHLIGHT_MAJOR,
    ALERT_HIGHLIGHT_COLORS,
    FONT_PRIMARY, SIZE_TITLE, SIZE_SUBTITLE, SIZE_BODY,
    SIZE_BODY_SMALL, SIZE_FOOTER, SIZE_NO_WARNING, SIZE_LABEL, SIZE_TINY,
    MAP_DAY1_WIDTH, MAP_DAY1_HEIGHT, MAP_SMALL_WIDTH, MAP_SMALL_HEIGHT,
    MARGIN_LEFT, MARGIN_RIGHT, PAGE_WIDTH,
)
from .table_helpers import (
    set_cell_shading, set_cell_borders, remove_all_borders,
    set_cell_vertical_alignment, set_cell_margins, set_row_height,
    set_cell_width, set_table_fixed_layout, set_table_col_widths,
    set_table_width, shade_run,
)
from .header_helpers import build_722e4_header, _add_run
from ..models.common import AlertLevel, HazardType
from ..models.seven22e4 import Seven22E4Bulletin, FiveDayEntry, HazardEntry
from ..i18n.translations import t

# Day name mapping
DAY_NAMES = {
    0: "Monday", 1: "Tuesday", 2: "Wednesday",
    3: "Thursday", 4: "Friday", 5: "Saturday", 6: "Sunday",
}

# Thin black border definition
THIN_BORDER = {'val': 'single', 'sz': 4, 'color': '000000'}


# Landscape A4: 29.7cm width - 2*1.27cm margins = 27.16cm = 15,398 twips
LANDSCAPE_TOTAL_W = 15398


class Seven22E4Builder(BaseBulletinBuilder):
    """Builds the 722E_4 Five Days Severe Weather Impact-Based Forecast."""

    def __init__(self, data: Seven22E4Bulletin, maps_dir: str = None):
        super().__init__(data, language="en", maps_dir=maps_dir)
        self.bulletin: Seven22E4Bulletin = data

    def _apply_orientation(self):
        """Switch to landscape for 722E_4 bulletins with tighter margins."""
        self._set_landscape()
        # Reference PDFs use tighter top/bottom margins in landscape
        section = self.doc.sections[0]
        section.top_margin = Cm(0.6)
        section.bottom_margin = Cm(0.5)

    def _format_date(self, d: date) -> str:
        return d.strftime("%d-%m-%Y")

    def _day_name(self, d: date) -> str:
        return DAY_NAMES[d.weekday()]

    def _get_issue_text(self) -> str:
        day_name = self._day_name(self.bulletin.issue_date)
        date_str = self._format_date(self.bulletin.issue_date)
        time_str = self.bulletin.issue_time.strftime("%H:%M")
        return f"Issued on {day_name}: {date_str}: {time_str} (EAT)"

    def _set_cell_borders_all(self, cell):
        """Apply thin black borders to all sides of a cell."""
        set_cell_borders(cell, top=THIN_BORDER, bottom=THIN_BORDER,
                        left=THIN_BORDER, right=THIN_BORDER)

    def _remove_spacing_between_tables(self, table1, table2):
        """Remove the paragraph gap between two consecutive tables."""
        tbl1 = table1._tbl
        tbl2 = table2._tbl
        body = tbl1.getparent()
        # Remove all paragraphs between the two tables
        to_remove = []
        next_el = tbl1.getnext()
        while next_el is not None and next_el != tbl2:
            if next_el.tag == qn('w:p'):
                to_remove.append(next_el)
            next_el = next_el.getnext()
        for el in to_remove:
            body.remove(el)

    def _merge_cells_in_row(self, table, row_idx, start_col, end_col):
        """Merge cells in a row from start_col to end_col (inclusive)."""
        cell_a = table.cell(row_idx, start_col)
        cell_b = table.cell(row_idx, end_col)
        cell_a.merge(cell_b)
        return table.cell(row_idx, start_col)

    def _alert_fill(self, level: AlertLevel) -> str:
        """Solid fill for alert label / Impacts expected chips (official TMA samples)."""
        key = getattr(level, "name", None) or str(level)
        if isinstance(level, AlertLevel):
            key = level.name
        return ALERT_HIGHLIGHT_COLORS.get(key, HEX_HIGHLIGHT_ADVISORY)

    def _shaded_label_run(self, paragraph, text: str, fill_hex: str, *,
                          bold=True, size=SIZE_BODY, underline=False):
        """Black text on solid yellow/orange/red fill — matches Downloads/tma samples."""
        run = paragraph.add_run(text)
        run.font.name = FONT_PRIMARY
        run.font.size = size
        run.font.color.rgb = COLOR_BLACK
        run.bold = bold
        run.underline = underline
        shade_run(run, fill_hex)
        return run

    def _build_day_content(self, cell, day: FiveDayEntry):
        """Build the forecast content for a single day inside a table cell.

        Official 722E_4 look (Downloads/tma samples):
          - ADVISORY / WARNING / MAJOR WARNING word is filled yellow / orange / red (black text)
          - "Impacts expected:" always present, filled with the same colour as the alert
          - Multiple hazards on one day each get their own coloured chips
        """
        # Clear default paragraph
        cell.paragraphs[0].clear()

        if day.is_no_warning:
            # Large "NO WARNING." text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            _add_run(p, t("no_warning", "en"), size=SIZE_NO_WARNING, bold=True)
            return

        hazards = list(day.hazards or [])
        if not hazards:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_run(p, t("no_warning", "en"), size=SIZE_NO_WARNING, bold=True)
            return

        # Build content for each hazard / scenario on this day
        for i, hazard in enumerate(hazards):
            fill = self._alert_fill(hazard.alert_level)
            if i > 0:
                # Visual separator between multiple scenarios on the same day
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)

            # Alert level (filled) + description on same line
            p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            alert_label = self._get_alert_label(hazard.alert_level)
            self._shaded_label_run(p, alert_label, fill, bold=True, size=SIZE_BODY)

            # Description text (not filled)
            if hazard.description:
                desc = hazard.description.strip()
                # Avoid doubling the alert word if description already starts with it
                for prefix in (alert_label, alert_label.title(), alert_label.lower()):
                    if desc.upper().startswith(prefix.upper()):
                        desc = desc[len(prefix):].lstrip(" :,-")
                        break
                if desc:
                    desc_run = p.add_run(" " + desc)
                    desc_run.font.name = FONT_PRIMARY
                    desc_run.font.size = SIZE_BODY
                    desc_run.font.color.rgb = COLOR_BLACK

            # Likelihood and Impact
            if hazard.rating:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)
                likelihood_val = t(f"level_{hazard.rating.likelihood.value.lower()}", "en")
                _add_run(p, f"{t('likelihood', 'en')}: {likelihood_val}", size=SIZE_BODY)

                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                impact_val = t(f"level_{hazard.rating.impact.value.lower()}", "en")
                _add_run(p, f"{t('impact', 'en')}: {impact_val}", size=SIZE_BODY)

            # "Impacts expected:" — ALWAYS present, filled with the alert colour (TMA sample)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            self._shaded_label_run(
                p, t("impacts_expected", "en"), fill, bold=True, size=SIZE_BODY,
            )

            body = (hazard.impacts_expected or "").strip()
            if body:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                _add_run(p, body, size=SIZE_BODY)

            # "Please be prepared." — bold
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, t("please_be_prepared", "en"), size=SIZE_BODY, bold=True)

    def _get_alert_label(self, level: AlertLevel) -> str:
        mapping = {
            AlertLevel.NO_WARNING: t("no_warning", "en"),
            AlertLevel.ADVISORY: t("advisory", "en"),
            AlertLevel.WARNING: t("warning", "en"),
            AlertLevel.MAJOR_WARNING: t("major_warning", "en"),
        }
        return mapping.get(level, "")

    # Icon file stems under assets/icons (same set the map_generator uses).
    _HAZARD_ICON_STEM = {
        HazardType.HEAVY_RAIN: "heavy_rain",
        HazardType.STRONG_WIND: "strong_wind",
        HazardType.LARGE_WAVES: "large_waves",
        HazardType.EXTREME_TEMPERATURE: "extreme_temperature",
        HazardType.FLOODS: "floods",
        HazardType.LANDSLIDES: "landslides",
    }
    # Translation key for the KEY strip label (falls back to hazard_type name).
    _HAZARD_KEY_I18N = {
        HazardType.HEAVY_RAIN: "heavy_rain_key",
        HazardType.STRONG_WIND: "strong_wind_key",
        HazardType.LARGE_WAVES: "large_waves_key",
        HazardType.EXTREME_TEMPERATURE: "extreme_temperature_key",
        HazardType.FLOODS: "hazard_floods",
        HazardType.LANDSLIDES: "hazard_landslides",
    }

    def _bulletin_hazard_types(self, day: FiveDayEntry = None) -> list:
        """Ordered unique hazard types present on a day (or the whole bulletin).

        The KEY must show the hazards that were actually authored / drawn — not a
        hard-coded Heavy rain icon when the operator issued wind or waves.
        """
        seen = []
        days = [day] if day is not None else list(self.bulletin.days)
        for d in days:
            if d is None:
                continue
            for h in (d.hazards or []):
                ht = getattr(h, "hazard_type", None)
                if ht is None:
                    continue
                if ht not in seen:
                    seen.append(ht)
        return seen

    def _hazard_icon_path(self, hazard_type: HazardType) -> Path | None:
        stem = self._HAZARD_ICON_STEM.get(hazard_type)
        if not stem:
            return None
        icons = Path(__file__).parent.parent.parent / "assets" / "icons"
        for name in (f"{stem}.png", f"{stem}_128.png", f"{stem}_64.png"):
            p = icons / name
            if p.exists():
                return p
        return None

    def _hazard_key_label(self, hazard_type: HazardType) -> str:
        i18n_key = self._HAZARD_KEY_I18N.get(hazard_type)
        if i18n_key:
            label = t(i18n_key, "en")
            # t() may return the key itself when missing — fall back to readable name.
            if label and label != i18n_key:
                return label
        name = getattr(hazard_type, "name", None) or str(hazard_type)
        return name.replace("_", " ").title()

    def _build_key_section(self, cell, day: FiveDayEntry = None):
        """Build the KEY section: one icon+label per hazard actually present, then colour chips.

        Uses in-paragraph shaded runs (not nested tables) so the Day-1 two-column
        map | forecast layout is not disturbed.
        """
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, t("key_label", "en"), size=SIZE_LABEL, bold=True)

        # Prefer hazards on the day whose map sits above this KEY; else whole bulletin.
        hazard_types = self._bulletin_hazard_types(day)
        if not hazard_types:
            hazard_types = self._bulletin_hazard_types(None)

        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)

        if not hazard_types:
            # NO WARNING day — still show a neutral KEY line (no false Heavy-rain claim).
            _add_run(p, " " + t("no_warning", "en"), size=SIZE_BODY)
        else:
            for ht in hazard_types:
                icon_path = self._hazard_icon_path(ht)
                if icon_path is not None and icon_path.exists():
                    run = p.add_run()
                    run.add_picture(str(icon_path), width=Cm(0.75), height=Cm(0.75))
                else:
                    run = p.add_run("\u26A0")
                    run.font.name = FONT_PRIMARY
                    run.font.size = Pt(12)
                _add_run(p, " " + self._hazard_key_label(ht) + "   ", size=SIZE_BODY)

        # Legend chips: yellow / orange / red boxes + labels (Downloads/tma KEY strip)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

        items = [
            (HEX_HIGHLIGHT_ADVISORY, t("advisory", "en")),
            (HEX_HIGHLIGHT_WARNING, t("warning", "en")),
            (HEX_HIGHLIGHT_MAJOR, t("major_warning", "en")),
        ]
        for i, (fill, label) in enumerate(items):
            # Wide shaded spaces act as the colour box
            chip = p.add_run("\u00a0\u00a0\u00a0\u00a0")
            chip.font.name = FONT_PRIMARY
            chip.font.size = Pt(12)
            shade_run(chip, fill)
            lab = p.add_run(f" {label}  ")
            lab.font.name = FONT_PRIMARY
            lab.font.size = Pt(8)
            lab.bold = True
            lab.font.color.rgb = COLOR_BLACK

    def _build_page1(self):
        """Build Page 1: Header + Issue Line + Day 1 table + Footer."""
        # Header (without issue line — it goes into the heading table instead)
        build_722e4_header(
            self.doc,
            include_issue_line=False,
            total_width_twips=LANDSCAPE_TOTAL_W,
        )

        day1 = self.bulletin.days[0]

        # Column widths in twips for landscape A4 (~27.16cm usable)
        total_w = LANDSCAPE_TOTAL_W
        col_left = total_w // 2    # ~13.58cm (map + KEY)
        col_right = total_w - col_left  # ~13.58cm (forecast content)

        # --- Day heading row: 1-column table ---
        heading_table = self.doc.add_table(rows=1, cols=1)
        heading_table.autofit = False
        heading_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(heading_table)
        set_table_width(heading_table, total_w)
        set_table_col_widths(heading_table, [total_w])

        heading_cell = heading_table.cell(0, 0)
        # Borders: top, left, right; no bottom (content table provides its own top)
        no_border = {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}
        set_cell_borders(heading_cell,
                        top=THIN_BORDER, left=THIN_BORDER, right=THIN_BORDER,
                        bottom=no_border)

        # Issue text — directly on top of the day heading, inside the table
        p = heading_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, self._get_issue_text(), size=SIZE_TITLE, bold=True,
                 color=COLOR_BLACK)

        # Day heading
        p = heading_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        day_name = self._day_name(day1.forecast_date)
        date_str = self._format_date(day1.forecast_date)
        _add_run(p, f"{day_name}: {date_str}", size=SIZE_SUBTITLE, bold=True)

        # --- Content row: 2-column table ---
        content_table = self.doc.add_table(rows=1, cols=2)
        content_table.autofit = False
        content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_fixed_layout(content_table)
        set_table_col_widths(content_table, [col_left, col_right])
        set_table_width(content_table, total_w)

        # Remove spacing between the two tables
        self._remove_spacing_between_tables(heading_table, content_table)

        # Left cell: Map + KEY
        left_cell = content_table.cell(0, 0)
        set_cell_width(left_cell, col_left)
        self._set_cell_borders_all(left_cell)

        p = left_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        self._add_image_or_placeholder(p, day1.map_image, MAP_DAY1_WIDTH, MAP_DAY1_HEIGHT)

        # KEY section below map — icons/labels for hazards actually on Day 1 (not hard-coded rain)
        self._build_key_section(left_cell, day=day1)

        # Right cell: Forecast content
        right_cell = content_table.cell(0, 1)
        set_cell_width(right_cell, col_right)
        self._set_cell_borders_all(right_cell)
        set_cell_vertical_alignment(right_cell, "top")
        self._build_day_content(right_cell, day1)

    def _build_page2(self):
        """Build Page 2: Header + Days 2-5 in a 4-column table.

        Layout: 3 rows x 4 columns
          Row 0: Day headings
          Row 1: Map images
          Row 2: Forecast text content
        """
        self._add_page_break()

        # Header (without issue line)
        build_722e4_header(self.doc, include_issue_line=False,
                           total_width_twips=LANDSCAPE_TOTAL_W)

        # 4-column table for Days 2-5 with 3 rows
        table = self.doc.add_table(rows=3, cols=4)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fixed layout: 4 equal columns across landscape width
        col_w = LANDSCAPE_TOTAL_W // 4  # ~6.79cm each
        set_table_fixed_layout(table)
        set_table_col_widths(table, [col_w, col_w, col_w, col_w])
        set_table_width(table, col_w * 4)

        for day_idx in range(4):
            day = self.bulletin.days[day_idx + 1]  # Days 2-5

            # --- Row 0: Day heading ---
            header_cell = table.cell(0, day_idx)
            set_cell_width(header_cell, col_w)
            self._set_cell_borders_all(header_cell)

            p = header_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            day_name = self._day_name(day.forecast_date)
            date_str = self._format_date(day.forecast_date)
            _add_run(p, f"{day_name}: {date_str}", size=SIZE_BODY_SMALL, bold=True)

            # --- Row 1: Map image ---
            map_cell = table.cell(1, day_idx)
            set_cell_width(map_cell, col_w)
            self._set_cell_borders_all(map_cell)

            p = map_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            self._add_image_or_placeholder(p, day.map_image,
                                           MAP_SMALL_WIDTH, MAP_SMALL_HEIGHT)

            # --- Row 2: Forecast text content ---
            text_cell = table.cell(2, day_idx)
            set_cell_width(text_cell, col_w)
            self._set_cell_borders_all(text_cell)
            set_cell_vertical_alignment(text_cell, "top")

            self._build_day_content_below_map(text_cell, day)

    def _build_day_content_below_map(self, cell, day: FiveDayEntry):
        """Build forecast content below a map in a day cell (Page 2 — same TMA chip styling)."""
        if day.is_no_warning or not day.hazards:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            _add_run(p, t("no_warning", "en"), size=SIZE_BODY_SMALL, bold=True)
            return

        for i, hazard in enumerate(day.hazards):
            fill = self._alert_fill(hazard.alert_level)
            if i > 0:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(0)

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            alert_label = self._get_alert_label(hazard.alert_level)
            self._shaded_label_run(
                p, alert_label, fill, bold=True, size=SIZE_BODY_SMALL,
            )

            if hazard.description:
                desc = hazard.description.strip()
                for prefix in (alert_label, alert_label.title(), alert_label.lower()):
                    if desc.upper().startswith(prefix.upper()):
                        desc = desc[len(prefix):].lstrip(" :,-")
                        break
                if desc:
                    desc_run = p.add_run(" " + desc)
                    desc_run.font.name = FONT_PRIMARY
                    desc_run.font.size = SIZE_BODY_SMALL
                    desc_run.font.color.rgb = COLOR_BLACK

            if hazard.rating:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(0)
                likelihood_val = t(f"level_{hazard.rating.likelihood.value.lower()}", "en")
                _add_run(p, f"{t('likelihood', 'en')}: {likelihood_val}", size=SIZE_BODY_SMALL)

                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                impact_val = t(f"level_{hazard.rating.impact.value.lower()}", "en")
                _add_run(p, f"{t('impact', 'en')}: {impact_val}", size=SIZE_BODY_SMALL)

            # Always show Impacts expected: with matching colour fill
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            self._shaded_label_run(
                p, t("impacts_expected", "en"), fill, bold=True, size=SIZE_BODY_SMALL,
            )

            body = (hazard.impacts_expected or "").strip()
            if body:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                _add_run(p, body, size=SIZE_BODY_SMALL)

            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            _add_run(p, t("please_be_prepared", "en"), size=SIZE_BODY_SMALL, bold=True)

    def _add_blue_line(self):
        """Add a solid blue horizontal line using paragraph bottom border."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Use paragraph border (bottom) to create a solid blue line
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt line
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '0070C0')  # TMA blue
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _add_footer(self):
        """Add the TMA footer with blue line and address."""
        self._add_blue_line()

        # Address block — centered
        footer_lines = [
            (t("correspondence", "en"), True),
            (t("director_general", "en"), False),
            (t("address_line1", "en"), False),
            (t("address_line2", "en"), False),
        ]
        for text, is_italic in footer_lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(10)
            run = p.add_run(text)
            run.font.name = FONT_PRIMARY
            run.font.size = SIZE_FOOTER
            run.italic = is_italic

        # Email line with blue links
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10)
        for part_text, part_color in [
            ("Email: ", None),
            ("met@meteo.go.tz", COLOR_BLUE_TMA),
            ("; Website: ", None),
            ("www.meteo.go.tz", COLOR_BLUE_TMA),
        ]:
            run = p.add_run(part_text)
            run.font.name = FONT_PRIMARY
            run.font.size = SIZE_FOOTER
            if part_color:
                run.font.color.rgb = part_color

        # ISO certification — centered, bold
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(t("iso_cert", "en"))
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_FOOTER
        run.bold = True

        # Form number — right aligned, tiny
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(t("form_number", "en"))
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_TINY

    def build(self) -> Document:
        """Build the complete 722E_4 document."""
        self._build_page1()
        self._add_footer()
        self._build_page2()
        self._add_footer()
        return self.doc
